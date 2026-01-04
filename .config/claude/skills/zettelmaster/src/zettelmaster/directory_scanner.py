#!/usr/bin/env python3
"""
Directory Scanner - Mechanical directory processing for batch analysis.
Reads entire directories as cohesive units without semantic analysis.
"""

import os
import base64
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
import json


@dataclass
class DirectoryContent:
    """Structure representing a directory's content."""
    path: Path
    text_files: Dict[str, str] = field(default_factory=dict)  # path -> content
    images: Dict[str, str] = field(default_factory=dict)  # path -> path or base64
    subdirectories: Dict[str, 'DirectoryContent'] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)


# Optional imports for document extraction
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    import io
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class DirectoryScanner:
    """Scans and reads directory content for batch processing."""
    
    # Text file extensions to process
    TEXT_EXTENSIONS = {'.md', '.txt', '.rst', '.org', '.tex', '.adoc', '.html', '.htm'}
    
    # Image file extensions to track
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.svg', '.webp', '.bmp'}
    
    # Document extensions that need special handling
    DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.odt'}
    
    def __init__(self, base_path: Path):
        """Initialize scanner with base path."""
        self.base_path = Path(base_path)
        
    def scan_directory(self, 
                      dir_path: Optional[Path] = None,
                      recursive: bool = True,
                      include_images: bool = True) -> DirectoryContent:
        """
        Scan directory and return structured content.
        
        Args:
            dir_path: Directory to scan (defaults to base_path)
            recursive: Whether to scan subdirectories
            include_images: Whether to track image files
            
        Returns:
            DirectoryContent with all files and structure
        """
        if dir_path is None:
            dir_path = self.base_path
        else:
            dir_path = Path(dir_path)
            
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")
            
        content = DirectoryContent(path=dir_path)
        
        # Scan directory
        for item in dir_path.iterdir():
            if item.is_file():
                # Process based on file type
                suffix = item.suffix.lower()
                
                if suffix in self.TEXT_EXTENSIONS:
                    # Read text content
                    try:
                        with open(item, 'r', encoding='utf-8') as f:
                            rel_path = str(item.relative_to(self.base_path))
                            content.text_files[rel_path] = f.read()
                    except Exception as e:
                        print(f"Error reading {item}: {e}")
                        
                elif include_images and suffix in self.IMAGE_EXTENSIONS:
                    # Store image path (not reading binary to save memory)
                    rel_path = str(item.relative_to(self.base_path))
                    content.images[rel_path] = str(item.absolute())
                    
                elif suffix in self.DOCUMENT_EXTENSIONS:
                    # Note: Special handling for PDFs/docs would go here
                    # For now, just note in metadata
                    rel_path = str(item.relative_to(self.base_path))
                    content.metadata.setdefault('documents', []).append(rel_path)
                    
            elif item.is_dir() and recursive:
                # Skip hidden directories
                if not item.name.startswith('.'):
                    # Recursively scan subdirectory
                    subdir_content = self.scan_directory(item, recursive=True, include_images=include_images)
                    if subdir_content.text_files or subdir_content.images or subdir_content.subdirectories:
                        content.subdirectories[item.name] = subdir_content
                        
        # Add directory metadata
        content.metadata['total_text_files'] = len(content.text_files)
        content.metadata['total_images'] = len(content.images)
        content.metadata['total_subdirs'] = len(content.subdirectories)
        content.metadata['directory_name'] = dir_path.name
        
        return content

    def extract_document_content(
        self,
        file_path: Path,
        enable_ocr: bool = False,
        max_pages: Optional[int] = None
    ) -> Optional[str]:
        """
        Extract text content from PDF or DOCX files.
        
        Args:
            file_path: Path to the document file
            enable_ocr: Enable OCR for scanned PDFs (requires pytesseract)
            max_pages: Maximum pages to extract from PDF (None = all)
        
        Returns:
            Extracted text or None if extraction failed
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf_content(file_path, enable_ocr, max_pages)
        elif suffix in ['.docx', '.odt']:
            return self._extract_docx_content(file_path)
        elif suffix in ['.doc', '.rtf']:
            # Legacy formats - would need python-docx2txt or similar
            print(f"Legacy format {suffix} not yet supported: {file_path.name}")
            return None
        else:
            return None
    
    def _extract_pdf_content(
        self,
        pdf_path: Path,
        enable_ocr: bool = False,
        max_pages: Optional[int] = None
    ) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            print("PyPDF2 not installed. Install with: pip install PyPDF2")
            return None
        
        try:
            extracted_text = []
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                num_pages = len(pdf_reader.pages)
                
                if max_pages:
                    num_pages = min(num_pages, max_pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    # If no text and OCR is enabled, try OCR
                    if not text.strip() and enable_ocr and OCR_AVAILABLE:
                        text = self._ocr_pdf_page(page)
                    
                    if text:
                        extracted_text.append(f"--- Page {page_num + 1} ---")
                        extracted_text.append(text)
            
            return '\n'.join(extracted_text)
            
        except Exception as e:
            print(f"Error extracting PDF {pdf_path.name}: {e}")
            return None
    
    def _extract_docx_content(self, docx_path: Path) -> Optional[str]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            print("python-docx not installed. Install with: pip install python-docx")
            return None
        
        try:
            doc = DocxDocument(docx_path)
            paragraphs = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            print(f"Error extracting DOCX {docx_path.name}: {e}")
            return None
    
    def _ocr_pdf_page(self, page) -> str:
        """OCR a PDF page (requires pytesseract and PIL)."""
        if not OCR_AVAILABLE:
            return ""
        
        try:
            # This is a simplified version - would need pdf2image for proper implementation
            # For now, return empty as full OCR implementation would be complex
            return ""
        except Exception as e:
            print(f"OCR failed: {e}")
            return ""
    
    def scan_directory_enhanced(
        self,
        dir_path: Optional[Path] = None,
        recursive: bool = True,
        include_images: bool = True,
        extract_documents: bool = True,
        enable_ocr: bool = False,
        max_pdf_pages: Optional[int] = None
    ) -> DirectoryContent:
        """
        Enhanced directory scanning with document extraction.
        
        Args:
            dir_path: Directory to scan (defaults to base_path)
            recursive: Whether to scan subdirectories
            include_images: Whether to track image files
            extract_documents: Whether to extract content from PDFs/DOCX
            enable_ocr: Enable OCR for scanned PDFs
            max_pdf_pages: Max pages to extract from PDFs
        
        Returns:
            DirectoryContent with all files and extracted content
        """
        if dir_path is None:
            dir_path = self.base_path
        else:
            dir_path = Path(dir_path)
        
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")
        
        content = DirectoryContent(path=dir_path)
        
        # Track document extraction stats
        extraction_stats = {
            'pdfs_extracted': 0,
            'docx_extracted': 0,
            'extraction_failed': 0
        }
        
        for item in dir_path.iterdir():
            if item.is_file():
                suffix = item.suffix.lower()
                rel_path = str(item.relative_to(self.base_path))
                
                if suffix in self.TEXT_EXTENSIONS:
                    # Read text content as before
                    try:
                        with open(item, 'r', encoding='utf-8') as f:
                            content.text_files[rel_path] = f.read()
                    except Exception as e:
                        print(f"Error reading {item}: {e}")
                
                elif include_images and suffix in self.IMAGE_EXTENSIONS:
                    # Store image path
                    content.images[rel_path] = str(item.absolute())
                
                elif suffix in self.DOCUMENT_EXTENSIONS and extract_documents:
                    # Extract document content
                    extracted = self.extract_document_content(
                        item,
                        enable_ocr=enable_ocr,
                        max_pages=max_pdf_pages
                    )
                    
                    if extracted:
                        # Store extracted text with .extracted suffix
                        extracted_key = f"{rel_path}.extracted"
                        content.text_files[extracted_key] = extracted
                        
                        if suffix == '.pdf':
                            extraction_stats['pdfs_extracted'] += 1
                        elif suffix in ['.docx', '.odt']:
                            extraction_stats['docx_extracted'] += 1
                    else:
                        extraction_stats['extraction_failed'] += 1
                        content.metadata.setdefault('failed_extractions', []).append(rel_path)
                
                elif suffix in self.DOCUMENT_EXTENSIONS:
                    # Just note in metadata if not extracting
                    content.metadata.setdefault('documents', []).append(rel_path)
            
            elif item.is_dir() and recursive:
                if not item.name.startswith('.'):
                    subdir_content = self.scan_directory_enhanced(
                        item,
                        recursive=True,
                        include_images=include_images,
                        extract_documents=extract_documents,
                        enable_ocr=enable_ocr,
                        max_pdf_pages=max_pdf_pages
                    )
                    if subdir_content.text_files or subdir_content.images or subdir_content.subdirectories:
                        content.subdirectories[item.name] = subdir_content
        
        # Add metadata
        content.metadata['total_text_files'] = len(content.text_files)
        content.metadata['total_images'] = len(content.images)
        content.metadata['total_subdirs'] = len(content.subdirectories)
        content.metadata['directory_name'] = dir_path.name
        
        if extract_documents:
            content.metadata['extraction_stats'] = extraction_stats
        
        return content
    
    def get_flat_content(self, content: DirectoryContent) -> Dict[str, Any]:
        """
        Flatten directory content for easier processing.
        
        Returns dict with all text and images in flat structure.
        """
        flat = {
            'base_path': str(content.path),
            'all_text': {},
            'all_images': {},
            'structure': self._get_structure_tree(content)
        }
        
        def flatten_recursive(dc: DirectoryContent, prefix: str = ""):
            # Add text files
            for path, text in dc.text_files.items():
                flat['all_text'][path] = text
                
            # Add images
            for path, img_path in dc.images.items():
                flat['all_images'][path] = img_path
                
            # Process subdirectories
            for subdir_name, subdir_content in dc.subdirectories.items():
                flatten_recursive(subdir_content)
                
        flatten_recursive(content)
        return flat
    
    def _get_structure_tree(self, content: DirectoryContent, indent: int = 0) -> str:
        """
        Create a tree representation of directory structure.
        """
        tree = []
        indent_str = "  " * indent
        
        # Add current directory
        tree.append(f"{indent_str}{content.path.name}/")
        
        # Add files
        for text_file in sorted(content.text_files.keys()):
            name = Path(text_file).name
            tree.append(f"{indent_str}  📄 {name}")
            
        for image_file in sorted(content.images.keys()):
            name = Path(image_file).name
            tree.append(f"{indent_str}  🖼️ {name}")
            
        # Add subdirectories
        for subdir_name in sorted(content.subdirectories.keys()):
            subdir_tree = self._get_structure_tree(
                content.subdirectories[subdir_name], 
                indent + 1
            )
            tree.append(subdir_tree)
            
        return "\n".join(tree)
    
    def export_to_toon(self, content: DirectoryContent) -> str:
        """
        Export directory content to TOON format for LLM processing.
        """
        lines = []
        
        def add_content(dc: DirectoryContent, indent: int = 0):
            prefix = "  " * indent
            
            # Add directory info
            lines.append(f"{prefix}directory: {dc.path.name}")
            
            # Add text files
            if dc.text_files:
                lines.append(f"{prefix}  text_files")
                for path in sorted(dc.text_files.keys()):
                    name = Path(path).name
                    word_count = len(dc.text_files[path].split())
                    lines.append(f"{prefix}    {name}: {word_count} words")
                    
            # Add images
            if dc.images:
                lines.append(f"{prefix}  images")
                for path in sorted(dc.images.keys()):
                    name = Path(path).name
                    lines.append(f"{prefix}    {name}")
                    
            # Add subdirectories
            if dc.subdirectories:
                lines.append(f"{prefix}  subdirectories")
                for subdir_name in sorted(dc.subdirectories.keys()):
                    add_content(dc.subdirectories[subdir_name], indent + 2)
                    
        add_content(content)
        return "\n".join(lines)
    
    def get_reading_order(self, content: DirectoryContent) -> List[str]:
        """
        Suggest reading order based on common patterns.
        Returns ordered list of file paths.
        """
        order = []
        
        def process_dir(dc: DirectoryContent):
            files = list(dc.text_files.keys())
            
            # Common ordering patterns
            priority_patterns = [
                'readme', 'index', 'introduction', 'overview',
                'chapter', 'section', 'part'
            ]
            
            # Sort by priority patterns
            def sort_key(filepath):
                name = Path(filepath).name.lower()
                for i, pattern in enumerate(priority_patterns):
                    if pattern in name:
                        # Extract number if present
                        import re
                        match = re.search(r'(\d+)', name)
                        num = int(match.group(1)) if match else 0
                        return (i, num, name)
                return (len(priority_patterns), 0, name)
                
            files.sort(key=sort_key)
            order.extend(files)
            
            # Process subdirectories in alphabetical order
            for subdir_name in sorted(dc.subdirectories.keys()):
                process_dir(dc.subdirectories[subdir_name])
                
        process_dir(content)
        return order


def main():
    """Example usage."""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: directory_scanner.py <directory_path>")
        sys.exit(1)
        
    dir_path = Path(sys.argv[1])
    scanner = DirectoryScanner(dir_path.parent)
    
    print(f"Scanning directory: {dir_path}")
    content = scanner.scan_directory(dir_path)
    
    # Display structure
    print("\nDirectory Structure:")
    print(scanner._get_structure_tree(content))
    
    # Display TOON format
    print("\nTOON Format:")
    print(scanner.export_to_toon(content))
    
    # Display suggested reading order
    print("\nSuggested Reading Order:")
    for i, filepath in enumerate(scanner.get_reading_order(content), 1):
        print(f"{i}. {filepath}")


if __name__ == "__main__":
    main()